# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════╗
║         EXPORT PDF PAR RÉGION  —  Version Python                    ║
║         Auteur  : adapté pour mg.kouame                             ║
║         Date    : Fév 2026                                          ║
║                                                                      ║
║  DÉPENDANCES (installer une seule fois) :                           ║
║      pip install openpyxl python-docx                               ║
║                                                                      ║
║  LANCEMENT :                                                         ║
║      Double-cliquer sur ce fichier  OU  python export_pdf_par_region.py
╚══════════════════════════════════════════════════════════════════════╝
"""



import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import re
from datetime import datetime

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import sys

# ═══════════════════════════════════════════════════════════════
#  PALETTE DE COULEURS (identique à la version VBA)
# ═══════════════════════════════════════════════════════════════
CLR_HEADER_BG  = RGBColor(0x17, 0x2F, 0x6F)   # Bleu marine foncé
CLR_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)   # Blanc
CLR_TITLE_BG   = RGBColor(0x3F, 0x7F, 0xE4)   # Bleu roi
CLR_ROW_ODD    = RGBColor(0xFF, 0xFF, 0xFF)   # Blanc
CLR_ROW_EVEN   = RGBColor(0xE7, 0xEF, 0xF6)   # Bleu très clair
CLR_BORDER     = RGBColor(0xB3, 0xBC, 0xC3)   # Gris-bleu
CLR_FOOTER_TXT = RGBColor(0x64, 0x64, 0x64)   # Gris moyen
CLR_SEP_LINE   = RGBColor(0x17, 0x2F, 0x6F)   # Bleu marine (séparateur)

MAX_COL_NAME_LENGTH = 20  # Longueur max autorisée pour un nom de colonne


# ═══════════════════════════════════════════════════════════════
#  UTILITAIRES WORD / DOCX
# ═══════════════════════════════════════════════════════════════

def set_cell_bg(cell, rgb: RGBColor):
    """Applique une couleur de fond à une cellule Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_borders(cell, color: RGBColor, size: int = 4):
    """Ajoute des bordures sur les 4 côtés d'une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), hex_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Définit les marges internes d'une cellule (en twentieths of a point)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def set_row_height(row, height_cm: float):
    """Fixe la hauteur d'une ligne de tableau."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))  # 1 cm ≈ 567 twips
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)


def add_paragraph_border_top(paragraph, color: RGBColor, size: int = 6):
    """Ajoute une bordure supérieure à un paragraphe (séparateur pied de page)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    str(size))
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(top)
    pPr.append(pBdr)


def add_page_number_field(paragraph):
    """Insère le champ PAGE et NUMPAGES dans un paragraphe."""
    def _field(instr):
        run  = OxmlElement('w:r')
        fldB = OxmlElement('w:fldChar')
        fldB.set(qn('w:fldCharType'), 'begin')
        run.append(fldB)

        runI  = OxmlElement('w:r')
        instr_el = OxmlElement('w:instrText')
        instr_el.set(qn('xml:space'), 'preserve')
        instr_el.text = f' {instr} '
        runI.append(instr_el)

        runS  = OxmlElement('w:r')
        fldS  = OxmlElement('w:fldChar')
        fldS.set(qn('w:fldCharType'), 'separate')
        runS.append(fldS)

        runE  = OxmlElement('w:r')
        fldE  = OxmlElement('w:fldChar')
        fldE.set(qn('w:fldCharType'), 'end')
        runE.append(fldE)

        return run, runI, runS, runE

    p = paragraph._p
    for el in _field('PAGE'):
        p.append(el)

    slash = OxmlElement('w:r')
    t     = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = ' / '
    slash.append(t)
    p.append(slash)

    for el in _field('NUMPAGES'):
        p.append(el)


def add_tab_stop_right(paragraph, position_cm: float):
    """Ajoute un tab stop à droite dans un paragraphe."""
    pPr  = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'),  'right')
    tab.set(qn('w:pos'),  str(int(position_cm * 567)))
    tabs.append(tab)
    pPr.append(tabs)


# ═══════════════════════════════════════════════════════════════
#  GÉNÉRATION DU DOCUMENT WORD (.docx) PAR RÉGION
# ═══════════════════════════════════════════════════════════════

def build_docx(region_name: str, headers: list, rows: list) -> Document:
    """
    Crée un Document Word stylisé pour une région donnée.
    - region_name : nom de la région (utilisé pour le titre et le pied de page)
    - headers     : liste des noms de colonnes
    - rows        : liste de listes de valeurs (une par ligne de données)
    """
    doc = Document()

    # ── Mise en page : Paysage, marges réduites ────────────────
    section = doc.sections[0]
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    # En paysage on échange largeur et hauteur (A4)
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin    = Cm(1.3)
    section.bottom_margin = Cm(1.6)
    section.left_margin   = Cm(1.3)
    section.right_margin  = Cm(1.3)

    # ── Nombre de colonnes ─────────────────────────────────────
    n_cols = len(headers)

    # ── Créer le tableau : 1 ligne titre + 1 ligne en-têtes + N lignes données
    n_rows = 1 + 1 + len(rows)
    table  = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = 'Table Grid'

    # Largeur totale disponible en paysage A4 avec marges 1.3cm
    total_width_cm = 29.7 - 2 * 1.3  # ≈ 27.1 cm
    col_width_cm   = total_width_cm / n_cols

    # ── LIGNE 0 : Titre de région (fusionnée) ─────────────────
    title_row = table.rows[0]
    set_row_height(title_row, 1.0)

    # Fusionner toutes les cellules de la ligne titre
    title_cell = title_row.cells[0]
    if n_cols > 1:
        title_cell = title_row.cells[0].merge(title_row.cells[n_cols - 1])

    set_cell_bg(title_cell, CLR_TITLE_BG)
    set_cell_margins(title_cell, top=60, bottom=60, left=200, right=200)

    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run  = title_para.add_run(f"RÉGION : {region_name.upper()}")
    title_run.font.name  = 'Calibri'
    title_run.font.size  = Pt(14)
    title_run.font.bold  = True
    title_run.font.color.rgb = CLR_HEADER_FG

    # ── LIGNE 1 : En-têtes de colonnes ────────────────────────
    header_row = table.rows[1]
    set_row_height(header_row, 0.8)

    for col_idx, col_name in enumerate(headers):
        cell = header_row.cells[col_idx]
        set_cell_bg(cell, CLR_HEADER_BG)
        set_cell_borders(cell, CLR_BORDER, size=4)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(str(col_name))
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = CLR_HEADER_FG

    # ── LIGNES 2+ : Données avec alternance de couleurs ───────
    for row_idx, row_data in enumerate(rows):
        data_row = table.rows[2 + row_idx]
        set_row_height(data_row, 0.65)
        bg_color = CLR_ROW_EVEN if row_idx % 2 == 0 else CLR_ROW_ODD

        for col_idx, value in enumerate(row_data):
            cell = data_row.cells[col_idx]
            set_cell_bg(cell, bg_color)
            set_cell_borders(cell, CLR_BORDER, size=3)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run  = para.add_run(str(value) if value is not None else '')
            run.font.name  = 'Calibri'
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

    # ── Largeur des colonnes ───────────────────────────────────
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            cell.width = Cm(col_width_cm)

    # ── Répéter ligne d'en-tête sur chaque page ────────────────
    # Ligne 1 (en-têtes colonnes) = répétée
    hdr_tr   = table.rows[1]._tr
    hdr_trPr = hdr_tr.get_or_add_trPr()
    tblHdr   = OxmlElement('w:tblHeader')
    hdr_trPr.append(tblHdr)

    # ── PIED DE PAGE ───────────────────────────────────────────
    footer   = section.footer
    footer_p = footer.paragraphs[0]

    # Ligne de séparation bleue au-dessus
    add_paragraph_border_top(footer_p, CLR_SEP_LINE, size=6)

    # Tab stop à droite (25 cm) pour aligner "Page X / Y"
    add_tab_stop_right(footer_p, 25.0)

    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Texte gauche : Région + date
    run_left = footer_p.add_run(
        f"Région : {region_name}   |   Exporté le : {datetime.now().strftime('%d/%m/%Y')}"
    )
    run_left.font.name  = 'Calibri'
    run_left.font.size  = Pt(8)
    run_left.font.color.rgb = CLR_FOOTER_TXT

    # Tabulation → droite
    tab_run = footer_p.add_run('\t')
    tab_run.font.size = Pt(8)

    # "Page "
    run_page = footer_p.add_run('Page ')
    run_page.font.name  = 'Calibri'
    run_page.font.size  = Pt(8)
    run_page.font.color.rgb = CLR_FOOTER_TXT

    # Champs PAGE / NUMPAGES
    add_page_number_field(footer_p)

    # Style de police pour les champs
    for run in footer_p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(8)

    return doc


# ═══════════════════════════════════════════════════════════════
#  CONVERSION DOCX → PDF  (via LibreOffice ou MS Word)
# ═══════════════════════════════════════════════════════════════

def convert_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Convertit un .docx en .pdf.
    Essaie d'abord LibreOffice, puis MS Word via COM (Windows).
    Retourne le chemin du PDF produit.
    """
    pdf_path = os.path.join(output_dir,
                            os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')

    # ── Tentative 1 : LibreOffice ──────────────────────────────
    lo_candidates = [
        'libreoffice', 'soffice',
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ]
    for lo in lo_candidates:
        try:
            result = subprocess.run(
                [lo, '--headless', '--convert-to', 'pdf',
                 '--outdir', output_dir, docx_path],
                capture_output=True, timeout=60
            )
            if result.returncode == 0 and os.path.exists(pdf_path):
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # ── Tentative 2 : MS Word via COM (Windows uniquement) ────
    try:
        import comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc  = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # ── Tentative 3 : win32com (autre lib Windows) ────────────
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc  = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        word.Quit()
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # Si aucune conversion n'a fonctionné → on garde le .docx
    return docx_path


# ═══════════════════════════════════════════════════════════════
#  LOGIQUE PRINCIPALE D'EXPORT
# ═══════════════════════════════════════════════════════════════

def run_export(excel_path: str, sheet_name: str,
               group_col: str, output_dir: str,
               excluded_cols: list,
               progress_callback=None,
               log_callback=None):
    """
    Lit le fichier Excel, regroupe par `group_col`,
    génère un PDF par valeur unique.
    """
    def log(msg):
        if log_callback:
            log_callback(msg)

    def progress(val):
        if progress_callback:
            progress_callback(val)

    # ── Charger le fichier Excel ───────────────────────────────
    log(f"📂 Chargement : {excel_path}")
    wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
    ws = wb[sheet_name]

    rows_iter = ws.iter_rows(values_only=True)
    raw_headers = list(next(rows_iter))

    # ── Colonnes actives (exclure les colonnes désignées) ──────
    active_indices = [
        i for i, h in enumerate(raw_headers)
        if h is not None and str(h).strip() not in excluded_cols
    ]
    headers = [str(raw_headers[i]).strip() for i in active_indices]

    # ── Index de la colonne de regroupement ───────────────────
    try:
        group_idx_in_active = headers.index(group_col)
        group_idx_raw       = active_indices[group_idx_in_active]
    except ValueError:
        # Chercher parmi tous les en-têtes bruts
        group_idx_raw = next(
            (i for i, h in enumerate(raw_headers)
             if h and str(h).strip() == group_col), None
        )
        if group_idx_raw is None:
            raise ValueError(f"Colonne de regroupement '{group_col}' introuvable.")

    # ── Lire toutes les lignes de données ─────────────────────
    all_data = []
    for row in rows_iter:
        if all(v is None for v in row):
            continue
        all_data.append(row)

    log(f"✅ {len(all_data)} lignes lues — {len(headers)} colonnes actives")

    # ── Regrouper par valeur de la colonne choisie ─────────────
    groups = {}
    for row in all_data:
        key = str(row[group_idx_raw]).strip() if row[group_idx_raw] is not None else '(Sans valeur)'
        groups.setdefault(key, []).append(row)

    log(f"🗂️  {len(groups)} groupe(s) détecté(s) : {', '.join(sorted(groups.keys()))}")

    # ── Créer le dossier de sortie ─────────────────────────────
    os.makedirs(output_dir, exist_ok=True)

    total  = len(groups)
    errors = []

    for idx, (region_name, region_rows) in enumerate(sorted(groups.items())):
        log(f"  ▶ Traitement : {region_name} ({len(region_rows)} ligne(s))...")

        # Extraire seulement les colonnes actives
        clean_rows = [
            [row[i] for i in active_indices]
            for row in region_rows
        ]

        # Construire le document Word
        doc = build_docx(region_name, headers, clean_rows)

        # Nom de fichier sécurisé
        safe_name = re.sub(r'[\/\\:*?"<>|]', '_', region_name).strip()
        if not safe_name:
            safe_name = 'Groupe_Inconnu'

        sheet_safe = re.sub(r'[\/\\:*?"<>|]', '_', sheet_name).strip()
        docx_name  = f"{sheet_safe}_{safe_name}.docx"
        docx_path  = os.path.join(output_dir, docx_name)

        doc.save(docx_path)

        # Convertir en PDF
        result_path = convert_to_pdf(docx_path, output_dir)

        if result_path.endswith('.pdf'):
            # Supprimer le .docx intermédiaire
            try:
                os.remove(docx_path)
            except Exception:
                pass
            log(f"  ✅ PDF généré : {os.path.basename(result_path)}")
        else:
            log(f"  ⚠️  PDF non disponible — fichier Word conservé : {docx_name}")
            errors.append(region_name)

        progress(int((idx + 1) / total * 100))

    wb.close()
    return errors


# ═══════════════════════════════════════════════════════════════
#  INTERFACE GRAPHIQUE (UserForm tkinter)
# ═══════════════════════════════════════════════════════════════

class ExportApp(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Export PDF par Région — mg.kouame")
        self.resizable(False, False)
        self.configure(bg='#F0F4F8')

        # Variables
        self.excel_path    = tk.StringVar()
        self.use_active    = tk.BooleanVar(value=False)
        self.sheet_var     = tk.StringVar()
        self.group_col_var = tk.StringVar()
        self.output_dir    = tk.StringVar()

        self._sheets        = []
        self._columns       = []
        self._excluded_vars = {}   # nom_colonne → BooleanVar

        self._build_ui()

    # ── Construction de l'interface ────────────────────────────
    def _build_ui(self):
        PAD = {'padx': 12, 'pady': 6}

        # ── Titre ──────────────────────────────────────────────
        header = tk.Frame(self, bg='#172F6F', height=60)
        header.pack(fill='x')
        tk.Label(
            header,
            text="  📊  Export PDF par Région",
            font=('Calibri', 15, 'bold'),
            bg='#172F6F', fg='white', anchor='w'
        ).pack(side='left', padx=16, pady=12)

        # ── Corps principal ────────────────────────────────────
        body = tk.Frame(self, bg='#F0F4F8')
        body.pack(fill='both', expand=True, padx=20, pady=12)

        # Section 1 : Source
        self._section(body, "1. Source du fichier Excel")

        src_frame = tk.Frame(body, bg='#F0F4F8')
        src_frame.pack(fill='x', pady=(0, 8))

        tk.Entry(src_frame, textvariable=self.excel_path,
                 width=52, font=('Calibri', 10),
                 state='readonly').pack(side='left', padx=(0, 6))
        tk.Button(src_frame, text="Parcourir…",
                  font=('Calibri', 10), bg='#3F7FE4', fg='white',
                  relief='flat', cursor='hand2',
                  command=self._browse_file).pack(side='left')

        # Section 2 : Feuille
        self._section(body, "2. Feuille à traiter")

        sheet_frame = tk.Frame(body, bg='#F0F4F8')
        sheet_frame.pack(fill='x', pady=(0, 8))

        self.sheet_cb = ttk.Combobox(sheet_frame, textvariable=self.sheet_var,
                                     font=('Calibri', 10), width=38, state='readonly')
        self.sheet_cb.pack(side='left')
        self.sheet_cb.bind('<<ComboboxSelected>>', self._on_sheet_selected)

        # Section 3 : Colonne de regroupement
        self._section(body, "3. Colonne de regroupement (ex : REGION, DISTRICT…)")

        grp_frame = tk.Frame(body, bg='#F0F4F8')
        grp_frame.pack(fill='x', pady=(0, 8))

        self.group_cb = ttk.Combobox(grp_frame, textvariable=self.group_col_var,
                                     font=('Calibri', 10), width=38, state='readonly')
        self.group_cb.pack(side='left')

        # Section 4 : Colonnes à exclure
        self._section(body, "4. Colonnes à exclure du PDF  (cocher pour exclure)")

        self.excl_frame = tk.Frame(body, bg='#FFFFFF',
                                   relief='solid', bd=1)
        self.excl_frame.pack(fill='x', pady=(0, 8), ipady=4)

        tk.Label(self.excl_frame,
                 text="  ← Charger un fichier pour voir les colonnes",
                 font=('Calibri', 9, 'italic'), bg='#FFFFFF',
                 fg='#888888').pack(anchor='w', padx=8)

        # Section 5 : Dossier de sortie
        self._section(body, "5. Dossier de sortie")

        out_frame = tk.Frame(body, bg='#F0F4F8')
        out_frame.pack(fill='x', pady=(0, 8))

        tk.Entry(out_frame, textvariable=self.output_dir,
                 width=52, font=('Calibri', 10)).pack(side='left', padx=(0, 6))
        tk.Button(out_frame, text="Parcourir…",
                  font=('Calibri', 10), bg='#3F7FE4', fg='white',
                  relief='flat', cursor='hand2',
                  command=self._browse_output).pack(side='left')

        # ── Séparateur ─────────────────────────────────────────
        ttk.Separator(body, orient='horizontal').pack(fill='x', pady=10)

        # ── Boutons ────────────────────────────────────────────
        btn_frame = tk.Frame(body, bg='#F0F4F8')
        btn_frame.pack(fill='x')

        tk.Button(
            btn_frame, text="✖  Annuler",
            font=('Calibri', 11), bg='#E0E0E0', fg='#333333',
            relief='flat', cursor='hand2', width=14,
            command=self.destroy
        ).pack(side='right', padx=(6, 0))

        tk.Button(
            btn_frame, text="🚀  Lancer l'export",
            font=('Calibri', 11, 'bold'), bg='#172F6F', fg='white',
            relief='flat', cursor='hand2', width=18,
            command=self._launch_export
        ).pack(side='right')

        # ── Journal de log ─────────────────────────────────────
        self._section(body, "Journal")

        log_frame = tk.Frame(body, bg='#F0F4F8')
        log_frame.pack(fill='both', expand=True, pady=(0, 8))

        self.log_box = tk.Text(log_frame, height=9, font=('Consolas', 9),
                               bg='#1E1E2E', fg='#A8D8A8',
                               relief='flat', state='disabled',
                               wrap='word')
        scrollbar = ttk.Scrollbar(log_frame, command=self.log_box.yview)
        self.log_box.configure(yscrollcommand=scrollbar.set)
        self.log_box.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        # ── Barre de progression ───────────────────────────────
        self.progress_var = tk.IntVar(value=0)
        self.progress_bar = ttk.Progressbar(body, variable=self.progress_var,
                                            maximum=100, length=500)
        self.progress_bar.pack(fill='x', pady=(0, 4))

        self.progress_lbl = tk.Label(body, text='', bg='#F0F4F8',
                                     font=('Calibri', 9), fg='#555555')
        self.progress_lbl.pack()

    def _section(self, parent, title: str):
        """Ajoute un label de section stylisé."""
        tk.Label(parent, text=title,
                 font=('Calibri', 10, 'bold'),
                 bg='#F0F4F8', fg='#172F6F',
                 anchor='w').pack(fill='x', pady=(8, 2))

    # ── Parcourir fichier source ───────────────────────────────
    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Sélectionner le fichier Excel",
            filetypes=[("Fichiers Excel", "*.xlsx *.xlsm *.xls"), ("Tous", "*.*")]
        )
        if not path:
            return
        self.excel_path.set(path)

        # Pré-remplir dossier de sortie
        if not self.output_dir.get():
            self.output_dir.set(os.path.dirname(path))

        self._load_sheets(path)

    def _load_sheets(self, path: str):
        """Charge la liste des feuilles du classeur."""
        try:
            wb = openpyxl.load_workbook(path, read_only=True)
            self._sheets = wb.sheetnames
            wb.close()
            self.sheet_cb['values'] = self._sheets
            if self._sheets:
                self.sheet_var.set(self._sheets[0])
                self._load_columns(path, self._sheets[0])
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible de lire le fichier :\n{e}")

    def _on_sheet_selected(self, _event=None):
        """Recharge les colonnes quand la feuille change."""
        path = self.excel_path.get()
        sheet = self.sheet_var.get()
        if path and sheet:
            self._load_columns(path, sheet)

    def _load_columns(self, path: str, sheet: str):
        """
        Lit la première ligne du fichier pour récupérer les noms de colonnes.
        Vérifie que les noms ne dépassent pas MAX_COL_NAME_LENGTH caractères.
        """
        try:
            wb = openpyxl.load_workbook(path, read_only=True)
            ws = wb[sheet]
            first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
            wb.close()
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible de lire la feuille :\n{e}")
            return

        # ── Vérification longueur des noms ────────────────────
        long_cols = [
            str(h).strip() for h in first_row
            if h is not None and len(str(h).strip()) > MAX_COL_NAME_LENGTH
        ]
        if long_cols:
            msg = (
                f"⚠️  Les colonnes suivantes ont un nom de plus de "
                f"{MAX_COL_NAME_LENGTH} caractères :\n\n"
                + '\n'.join(f"  • {c}  ({len(c)} car.)" for c in long_cols)
                + "\n\nVeuillez renommer ces colonnes dans Excel "
                  "(moins de 30 caractères) avant de continuer."
            )
            messagebox.showwarning("Noms de colonnes trop longs", msg)
            return

        self._columns = [
            str(h).strip() for h in first_row
            if h is not None and str(h).strip() != ''
        ]

        # Mettre à jour la combobox de regroupement
        self.group_cb['values'] = self._columns
        if self._columns:
            self.group_col_var.set(self._columns[0])

        # Mettre à jour les cases à cocher d'exclusion
        self._refresh_exclusion_checkboxes()

    def _refresh_exclusion_checkboxes(self):
        """Reconstruit les cases à cocher pour l'exclusion de colonnes."""
        for widget in self.excl_frame.winfo_children():
            widget.destroy()
        self._excluded_vars.clear()

        if not self._columns:
            tk.Label(self.excl_frame,
                     text="  ← Aucune colonne détectée",
                     font=('Calibri', 9, 'italic'),
                     bg='#FFFFFF', fg='#888888').pack(anchor='w', padx=8)
            return

        # Afficher en grille de 3 colonnes
        cols_per_row = 3
        for i, col in enumerate(self._columns):
            var = tk.BooleanVar(value=False)
            self._excluded_vars[col] = var
            row_f = i // cols_per_row
            col_f = i % cols_per_row
            cb = tk.Checkbutton(
                self.excl_frame, text=col, variable=var,
                font=('Calibri', 9), bg='#FFFFFF', anchor='w',
                cursor='hand2'
            )
            cb.grid(row=row_f, column=col_f, sticky='w', padx=10, pady=2)

    def _browse_output(self):
        folder = filedialog.askdirectory(title="Sélectionner le dossier de sortie")
        if folder:
            self.output_dir.set(folder)

    # ── Lancement de l'export ──────────────────────────────────
    def _launch_export(self):
        # Validations
        if not self.excel_path.get():
            messagebox.showwarning("Fichier manquant", "Veuillez sélectionner un fichier Excel.")
            return
        if not self.sheet_var.get():
            messagebox.showwarning("Feuille manquante", "Veuillez sélectionner une feuille.")
            return
        if not self.group_col_var.get():
            messagebox.showwarning("Colonne manquante",
                                   "Veuillez sélectionner la colonne de regroupement.")
            return
        if not self.output_dir.get():
            messagebox.showwarning("Dossier manquant",
                                   "Veuillez indiquer le dossier de sortie.")
            return

        excluded = [col for col, var in self._excluded_vars.items() if var.get()]

        self._log_clear()
        self.progress_var.set(0)
        self.progress_lbl.config(text='Export en cours…')

        try:
            errors = run_export(
                excel_path=self.excel_path.get(),
                sheet_name=self.sheet_var.get(),
                group_col=self.group_col_var.get(),
                output_dir=self.output_dir.get(),
                excluded_cols=excluded,
                progress_callback=self._update_progress,
                log_callback=self._log
            )

            if errors:
                msg = (f"Export terminé avec {len(errors)} avertissement(s).\n"
                       f"PDF non générés (Word conservé) :\n"
                       + '\n'.join(f"  • {e}" for e in errors)
                       + f"\n\n📁 Dossier : {self.output_dir.get()}")
                messagebox.showwarning("Export terminé avec avertissements", msg)
            else:
                messagebox.showinfo(
                    "✅ Export terminé",
                    f"Tous les PDFs ont été générés avec succès !\n\n"
                    f"📁 Dossier : {self.output_dir.get()}"
                )
            self.progress_lbl.config(text='✅ Terminé')

        except Exception as e:
            messagebox.showerror("Erreur", f"Une erreur est survenue :\n{e}")
            self.progress_lbl.config(text='❌ Erreur')

    # ── Helpers log / progression ──────────────────────────────
    def _log(self, message: str):
        self.log_box.config(state='normal')
        self.log_box.insert('end', message + '\n')
        self.log_box.see('end')
        self.log_box.config(state='disabled')
        self.update_idletasks()

    def _log_clear(self):
        self.log_box.config(state='normal')
        self.log_box.delete('1.0', 'end')
        self.log_box.config(state='disabled')

    def _update_progress(self, value: int):
        self.progress_var.set(value)
        self.progress_lbl.config(text=f'{value}%')
        self.update_idletasks()


# ═══════════════════════════════════════════════════════════════
#  POINT D'ENTRÉE
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    app = ExportApp()
    app.mainloop()
